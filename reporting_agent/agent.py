import os
from google.adk.agents.llm_agent import Agent
from google.adk.models.lite_llm import LiteLlm
from google.adk.tools import ToolContext
import pandas as pd
import ast
from docx import Document
from litellm import completion
import os
from dotenv import load_dotenv
load_dotenv()

doc = Document()

def path_for_output_files(session_id: str):
    cwd = os.getcwd()
    global output_path
    output_path = os.path.join(cwd,"OSINT_Files",session_id)
    if not os.path.exists(output_path):
        os.makedirs(output_path)
    return output_path

def sudomains_reporting(file_path: str):
    doc.add_heading("Subdomains identified from the OSINT Assessment",0)
    doc.add_paragraph("This is a list of subdomains identified for the root domain:")
    try:
        with open(file_path, "r") as f:
            subdomains = f.readlines()
            for subdomain in subdomains:
                doc.add_paragraph(subdomain, style='List Bullet')
    except:
        return

def emails_reporting(file_path: str):
    doc.add_section()
    doc.add_heading("Email Addresses identified from the OSINT Assessment",0)
    doc.add_paragraph("This is a list of emails identified for the root domain:")
    try:
        with open(file_path, "r") as f:
            emails = f.readlines()
            for email in emails:
                doc.add_paragraph(email, style='List Bullet')
    except:
        return

def breach_check_reporting(file_path: str):
    doc.add_section()
    doc.add_heading("Breach results for the Email Addresses",0)
    try:
        with open(file_path, "r") as f:
            breach_results = f.readlines()
            for result in breach_results:
                doc.add_paragraph(result, style='List Bullet')
    except:
        return

def bbot_reporting(file_path: str):
    doc.add_section()
    try:
        bbot_data = pd.read_csv(file_path)
    except:
        return
    event_types = bbot_data['Event type'].unique()
    doc.add_heading("Annex: Raw BBOT Event Data",0)
    for etype in event_types:
        if etype == "SCAN":
            continue
        temporary_df = bbot_data[bbot_data['Event type'] == etype].drop(columns=['Source Module', 'Scope Distance', 'Event Tags', 'Discovery Path'])
        doc.add_heading("Event type: "+etype,1)
        for i in range(len(temporary_df)):
            try:
                event_data = ast.literal_eval(temporary_df.iloc[i]['Event data'])
            except:
                event_data = str(temporary_df.iloc[i]['Event data'])
            if type(event_data) == str:
                doc.add_paragraph("Event Data: "+event_data, style='List Bullet')
            if type(event_data) == dict:
                doc.add_paragraph("Event:", style='List Bullet')
                for key in event_data:
                    content = f"{key}: {event_data[key]}"
                    doc.add_paragraph(content, style='List Bullet 2')

def executive_summary_llm(content: str):
    response = completion(
        model="groq/openai/gpt-oss-120b", 
        messages=[
           {
            "role": "system",
            "content": """
                    You are a senior cybersecurity consultant preparing an Executive Summary for a professional security assessment report.
                    
                    Your task:
                    - Analyze the provided OSINT-derived content and distill it into a clear, high-level Executive Summary.
                    - Focus on factual, evidence-based security insights derived strictly from the provided content.
                    - Do NOT invent findings, speculate, or introduce assumptions not supported by the input.
                    
                    Audience & tone:
                    - Written for non-technical executives and business stakeholders.
                    - Clear, concise, professional, and risk-focused.
                    - Avoid tool names, raw data dumps, or technical noise unless essential for understanding risk.
                    
                    Content guidelines:
                    - Summarize the overall security posture and key themes.
                    - Highlight the most critical risks, exposures, or weaknesses.
                    - Briefly explain potential business impact (e.g., data exposure, reputational risk, regulatory risk).
                    - Optionally mention high-level mitigation themes (no step-by-step remediation).
                    
                    Formatting (for direct insertion into a DOCX report):
                    - Use short paragraphs.
                    - Use bullet points only if they improve clarity.
                    - Do NOT include markdown, code blocks, tables, emojis, or headings other than plain text.
                    - Do NOT repeat the phrase “Executive Summary” in the output.
                    
                    Length:
                    - Target 2–4 well-structured sections such as Scope, Testing Methodology, Highlighted Findings, Risk Analysis, Recommendations, prioritizing clarity over verbosity.
                    """,
        },
            {"role": "user", "content": content}
       ]
    )
    return(response.choices[0].message['content'])

def add_exec_summary():
    doc_till_now = []
    for item in doc.paragraphs:
        doc_till_now.append(item.text)
    content = "".join(doc_till_now)
    exec_summary = executive_summary_llm(content)
    doc.add_section()
    doc.add_heading("Executive Summary",0)
    doc.add_paragraph(exec_summary)

def reporting_start_function(tool_context: ToolContext):
    session_id = getattr(tool_context.session, "id")
    print(
        "\n*************\n"
        "reporting_agent invoked\n"
        "*************\n"
    )
    project_folder = path_for_output_files(session_id)

    global doc
    doc = Document()
    
    bbot_raw_file_path = os.path.join(project_folder, "bbot_output.csv")
    subdomains_file_path = os.path.join(project_folder, "bbot_subdomains.txt")
    emails_file_path = os.path.join(project_folder, "bbot_emails.txt")
    breach_check_results_path = os.path.join(project_folder, "breach_check_scan.txt")
    
    sudomains_reporting(subdomains_file_path)
    emails_reporting(emails_file_path)
    breach_check_reporting(breach_check_results_path)
    bbot_reporting(bbot_raw_file_path)
    add_exec_summary()

    filename = session_id+"_OSINT_Report.docx"
    file_location = os.path.join(project_folder,filename)
    doc.save(file_location)
    print(f"*********\nReport saved to the following location: {project_folder}\n*********\nFilename: {filename}\n*********\n")
    return(f"File stored in location: {file_location}")

root_agent = Agent(
    model=LiteLlm(model="groq/openai/gpt-oss-120b", timeout=10000, api_key=os.getenv('GROQ_API_KEY')),
    # model='gemini-2.5-pro',
    name="reporting_agent",
    description="An agent that aggregates OSINT findings from multiple agents and generates a professional .docx report.",
    instruction=(
        """
        You are a robust reporting agent that aggregates findings from multiple OSINT agents (bbot, whois, breachcheck).

        Your Tasks:
        1. When invoked, call the `reporting_start_function` tool which returns a string that contains the location of where the generated file is stored
        2. Return confirmation with report filepath

        The `reporting_start_function` tool will:
        - Parse all input formats flexibly (JSON, lists, text, dicts)
        - Validate and sanitize data
        - Generate a professional report with all sections
        - Save the report to project directory 


        Do not attempt to create reports manually. Always use the `reporting_start_function` tool.
        Strictly adhere to these instructions. Do not deviate.
        """
    ),
    tools=[reporting_start_function]
)
